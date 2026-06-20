from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from datetime import date, datetime, time, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import HTTPException

from ..core.config import FRONTEND_DIR
from ..db import db_session, utc_now
from .disease_prediction import precautions as disease_precautions


RESPIRATORY_KEYWORDS = [
    "cough",
    "breathing",
    "respiratory",
    "lung",
    "lungs",
    "chest",
    "pneumonia",
    "asthma",
    "flu",
    "cold",
    "bronchitis",
]
MUSCULOSKELETAL_KEYWORDS = [
    "bone",
    "fracture",
    "joint",
    "leg",
    "hand",
    "spine",
    "knee",
    "shoulder",
    "ankle",
    "arm",
    "muscle",
    "orthopedic",
]
CARDIOVASCULAR_KEYWORDS = [
    "heart",
    "blood pressure",
    "hypertension",
    "chest pain",
    "cardiac",
]
DIGESTIVE_KEYWORDS = [
    "stomach",
    "abdomen",
    "digestive",
    "nausea",
    "vomit",
    "diarrhea",
]
SKIN_KEYWORDS = [
    "skin",
    "rash",
    "itch",
    "allergy",
    "dermat",
]
DENTAL_KEYWORDS = [
    "dental",
    "tooth",
    "teeth",
    "mouth",
    "jaw",
]
NEURO_KEYWORDS = [
    "head",
    "brain",
    "migraine",
    "seizure",
    "stroke",
    "nerve",
]


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _safe_json_load(value: str | None, default: Any) -> Any:
    if not value:
        return default
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return default


def _safe_json_dump(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, default=str)


def _date_bounds(start_date: date, end_date: date) -> tuple[str, str]:
    start_dt = datetime.combine(start_date, time.min, tzinfo=timezone.utc)
    end_dt = datetime.combine(end_date, time.max, tzinfo=timezone.utc)
    return start_dt.isoformat(), end_dt.isoformat()


def _format_display_date(value: str) -> str:
    if isinstance(value, date):
        return value.strftime("%d %b %Y %I:%M %p")
    try:
        parsed = datetime.fromisoformat(value)
    except ValueError:
        return value
    return parsed.strftime("%d %b %Y %I:%M %p")


def _format_date_only(value: str) -> str:
    if isinstance(value, date):
        return value.strftime("%d %b %Y")
    try:
        parsed = datetime.fromisoformat(value)
    except ValueError:
        return value
    return parsed.strftime("%d %b %Y")


def _display_calendar_date(value: Any) -> str:
    if isinstance(value, date):
        return value.strftime("%d %b %Y")
    if isinstance(value, datetime):
        return value.strftime("%d %b %Y")
    if isinstance(value, str):
        try:
            return date.fromisoformat(value[:10]).strftime("%d %b %Y")
        except ValueError:
            try:
                return datetime.fromisoformat(value).strftime("%d %b %Y")
            except ValueError:
                return value
    return str(value)


def _extract_keywords(text: str) -> list[str]:
    normalized = _normalize_text(text).lower()
    if not normalized:
        return ["General Health"]

    matches: list[str] = []

    def add_if(keyword_list: list[str], label: str) -> None:
        if any(keyword in normalized for keyword in keyword_list) and label not in matches:
            matches.append(label)

    add_if(RESPIRATORY_KEYWORDS, "Respiratory Issues")
    add_if(MUSCULOSKELETAL_KEYWORDS, "Musculoskeletal Concerns")
    add_if(CARDIOVASCULAR_KEYWORDS, "Cardiovascular Concerns")
    add_if(DIGESTIVE_KEYWORDS, "Digestive Health")
    add_if(SKIN_KEYWORDS, "Skin Conditions")
    add_if(DENTAL_KEYWORDS, "Dental / Oral Health")
    add_if(NEURO_KEYWORDS, "Neurological Concerns")

    if not matches:
        matches.append("General Health")

    return matches[:3]


def _risk_from_prediction(disease: str, confidence: float) -> str:
    disease_lower = disease.lower()
    high_risk_keywords = [
        "pneumonia",
        "stroke",
        "heart",
        "cancer",
        "fracture",
        "tuberculosis",
        "kidney",
        "appendicitis",
        "severe",
    ]
    medium_risk_keywords = [
        "asthma",
        "arthritis",
        "bronchitis",
        "allergy",
        "diabetes",
        "migraine",
        "infection",
    ]

    if any(keyword in disease_lower for keyword in high_risk_keywords):
        return "High"
    if confidence >= 0.68 or any(keyword in disease_lower for keyword in medium_risk_keywords):
        return "Medium"
    return "Low"


def _follow_up_from_risk(risk_level: str, concern: str) -> list[str]:
    actions = [
        "Monitor symptoms and note any changes in intensity, duration, or frequency.",
        "Share the report with a qualified healthcare professional for formal review.",
    ]

    if risk_level == "High":
        actions.insert(
            0,
            "Arrange prompt specialist follow-up, especially if symptoms are worsening or severe.",
        )
    elif concern == "Respiratory Issues":
        actions.insert(0, "Seek timely evaluation if breathing symptoms, fever, or cough persist.")
    elif concern == "Musculoskeletal Concerns":
        actions.insert(0, "Consider orthopedic follow-up if pain, swelling, or mobility issues continue.")
    elif concern == "Cardiovascular Concerns":
        actions.insert(0, "Track blood pressure or chest symptoms and seek care promptly if they worsen.")

    return actions[:4]


def _classify_topic(text: str) -> str:
    return _extract_keywords(text)[0]


def _keywords_for_text(text: str) -> list[str]:
    return _extract_keywords(text)


def log_user_activity(
    connection,
    user_id: int,
    activity_type: str,
    title: str,
    description: str,
    *,
    created_at: str | None = None,
    source_key: str | None = None,
) -> int:
    timestamp = created_at or utc_now()
    cursor = connection.execute(
        """
        INSERT INTO activities (user_id, activity_type, title, description, created_at)
        VALUES (?, ?, ?, ?, ?)
        """,
        (user_id, activity_type, title, description, timestamp),
    )
    activity_id = cursor.lastrowid
    activity_source = source_key or f"activity:{activity_id}"
    connection.execute(
        """
        INSERT OR IGNORE INTO user_activity (source_key, user_id, activity_type, title, description, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (activity_source, user_id, activity_type, title, description, timestamp),
    )
    return activity_id


def record_disease_prediction(
    connection,
    user_id: int,
    *,
    source_key: str,
    created_at: str,
    input_symptoms: list[str],
    predictions: list[dict],
) -> None:
    if not predictions:
        return

    top_prediction = predictions[0]
    confidence = float(top_prediction.get("confidence") or 0.0)
    disease = str(top_prediction.get("disease") or "Unknown").strip() or "Unknown"
    risk_level = _risk_from_prediction(disease, confidence)
    recommendations = top_prediction.get("precautions") or []

    connection.execute(
        """
        INSERT OR IGNORE INTO disease_predictions (
            source_key, user_id, created_at, input_symptoms_json,
            predicted_disease, confidence_score, risk_level, recommendations_json, predictions_json
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            source_key,
            user_id,
            created_at,
            _safe_json_dump(input_symptoms),
            disease,
            confidence,
            risk_level,
            _safe_json_dump(recommendations),
            _safe_json_dump(predictions),
        ),
    )


def record_xray_analysis(
    connection,
    user_id: int,
    *,
    source_key: str,
    created_at: str,
    file_name: str,
    result: dict,
) -> None:
    connection.execute(
        """
        INSERT OR IGNORE INTO xray_analysis (
            source_key, user_id, created_at, image_type, findings_json,
            abnormalities_json, severity, confidence, summary, recommendations_json, file_name
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            source_key,
            user_id,
            created_at,
            str(result.get("image_type") or "Unknown"),
            _safe_json_dump(result.get("findings") or []),
            _safe_json_dump(result.get("abnormalities") or []),
            str(result.get("severity") or "Medium"),
            str(result.get("confidence") or "Unavailable"),
            str(result.get("summary") or ""),
            _safe_json_dump(result.get("recommendations") or []),
            file_name,
        ),
    )


def record_chat_history(
    connection,
    user_id: int,
    *,
    source_key: str,
    created_at: str,
    conversation_id: int | None,
    question: str,
    answer: str,
) -> None:
    topics = _keywords_for_text(question + " " + answer)
    follow_up = _follow_up_from_risk("Medium", topics[0] if topics else "General Health")
    connection.execute(
        """
        INSERT OR IGNORE INTO chat_history (
            source_key, user_id, conversation_id, created_at, question, answer, topics_json, follow_up_json
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            source_key,
            user_id,
            conversation_id,
            created_at,
            question,
            answer,
            _safe_json_dump(topics),
            _safe_json_dump(follow_up),
        ),
    )


def _seed_legacy_user_activity(connection, user_id: int) -> None:
    activities = connection.execute(
        """
        SELECT id, user_id, activity_type, title, description, created_at
        FROM activities
        WHERE user_id = ?
        ORDER BY created_at ASC
        """,
        (user_id,),
    ).fetchall()

    for row in activities:
        connection.execute(
            """
            INSERT OR IGNORE INTO user_activity (source_key, user_id, activity_type, title, description, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                f"activity:{row['id']}",
                row["user_id"],
                row["activity_type"],
                row["title"],
                row["description"],
                row["created_at"],
            ),
        )


def _parse_prediction_activity(description: str) -> tuple[list[str], str]:
    match = re.search(r"Analyzed symptoms:\s*(.*?)\.\s*Top result:\s*(.*)$", description)
    if match:
        raw_symptoms = [item.strip() for item in match.group(1).split(",") if item.strip()]
        disease = match.group(2).strip() or "Unknown"
        return raw_symptoms, disease

    return [], "Unknown"


def _seed_legacy_predictions(connection, user_id: int) -> None:
    rows = connection.execute(
        """
        SELECT id, description, created_at
        FROM activities
        WHERE user_id = ? AND activity_type = 'prediction'
        ORDER BY created_at ASC
        """,
        (user_id,),
    ).fetchall()

    for row in rows:
        symptoms, disease = _parse_prediction_activity(row["description"])
        precautions = disease_precautions.get(disease, [])
        connection.execute(
            """
            INSERT OR IGNORE INTO disease_predictions (
                source_key, user_id, created_at, input_symptoms_json,
                predicted_disease, confidence_score, risk_level, recommendations_json, predictions_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                f"activity:{row['id']}",
                user_id,
                row["created_at"],
                _safe_json_dump(symptoms),
                disease,
                0.0,
                "Unknown",
                _safe_json_dump(precautions),
                _safe_json_dump([
                    {
                        "disease": disease,
                        "confidence": 0.0,
                        "description": None,
                        "precautions": precautions,
                    }
                ]),
            ),
        )


def _parse_xray_activity(description: str) -> tuple[str, str]:
    type_match = re.search(r"Type:\s*(.*?)\.\s*Severity:", description)
    severity_match = re.search(r"Severity:\s*(.*)$", description)
    image_type = type_match.group(1).strip() if type_match else "Unknown"
    severity = severity_match.group(1).strip().rstrip(".") if severity_match else "Medium"
    return image_type or "Unknown", severity or "Medium"


def _seed_legacy_xray(connection, user_id: int) -> None:
    rows = connection.execute(
        """
        SELECT id, title, description, created_at
        FROM activities
        WHERE user_id = ? AND activity_type = 'xray'
        ORDER BY created_at ASC
        """,
        (user_id,),
    ).fetchall()

    for row in rows:
        image_type, severity = _parse_xray_activity(row["description"])
        connection.execute(
            """
            INSERT OR IGNORE INTO xray_analysis (
                source_key, user_id, created_at, image_type, findings_json,
                abnormalities_json, severity, confidence, summary, recommendations_json, file_name
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                f"activity:{row['id']}",
                user_id,
                row["created_at"],
                image_type,
                _safe_json_dump(["Legacy X-ray record imported from activity history."]),
                _safe_json_dump([]),
                severity,
                "Unavailable",
                f"Legacy X-ray record imported from activity history for {image_type}. This is not a medical diagnosis.",
                _safe_json_dump([
                    "Review the original scan with a qualified healthcare professional.",
                ]),
                row["title"] or "legacy-xray",
            ),
        )


def _seed_legacy_chat_history(connection, user_id: int) -> None:
    conversations = connection.execute(
        "SELECT id FROM conversations WHERE user_id = ? ORDER BY created_at ASC",
        (user_id,),
    ).fetchall()

    for conversation in conversations:
        messages = connection.execute(
            """
            SELECT id, role, content, created_at
            FROM messages
            WHERE conversation_id = ? AND user_id = ?
            ORDER BY created_at ASC, id ASC
            """,
            (conversation["id"], user_id),
        ).fetchall()

        last_user_message: dict | None = None
        for message in messages:
            if message["role"] == "user":
                last_user_message = message
                continue

            if message["role"] != "assistant" or not last_user_message:
                continue

            source_key = f"message:{message['id']}"
            topics = _keywords_for_text(last_user_message["content"] + " " + message["content"])
            follow_up = _follow_up_from_risk("Medium", topics[0] if topics else "General Health")
            connection.execute(
                """
                INSERT OR IGNORE INTO chat_history (
                    source_key, user_id, conversation_id, created_at, question, answer, topics_json, follow_up_json
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    source_key,
                    user_id,
                    conversation["id"],
                    message["created_at"],
                    last_user_message["content"],
                    message["content"],
                    _safe_json_dump(topics),
                    _safe_json_dump(follow_up),
                ),
            )


def sync_reporting_tables(connection, user_id: int) -> None:
    _seed_legacy_user_activity(connection, user_id)
    _seed_legacy_predictions(connection, user_id)
    _seed_legacy_xray(connection, user_id)
    _seed_legacy_chat_history(connection, user_id)


def _fetch_user_activity(connection, user_id: int, start_iso: str, end_iso: str) -> list[dict]:
    return connection.execute(
        """
        SELECT activity_type, title, description, created_at
        FROM user_activity
        WHERE user_id = ? AND created_at BETWEEN ? AND ?
        ORDER BY created_at ASC, id ASC
        """,
        (user_id, start_iso, end_iso),
    ).fetchall()


def _fetch_disease_predictions(connection, user_id: int, start_iso: str, end_iso: str) -> list[dict]:
    rows = connection.execute(
        """
        SELECT created_at, input_symptoms_json, predicted_disease, confidence_score,
               risk_level, recommendations_json, predictions_json
        FROM disease_predictions
        WHERE user_id = ? AND created_at BETWEEN ? AND ?
        ORDER BY created_at ASC, id ASC
        """,
        (user_id, start_iso, end_iso),
    ).fetchall()

    items = []
    for row in rows:
        items.append(
            {
                "prediction_date": _format_display_date(row["created_at"]),
                "input_symptoms": _safe_json_load(row["input_symptoms_json"], []),
                "predicted_disease": row["predicted_disease"],
                "confidence_score": f"{float(row['confidence_score']) * 100:.0f}%",
                "risk_level": row["risk_level"],
                "recommendations": _safe_json_load(row["recommendations_json"], []),
                "_confidence_raw": float(row["confidence_score"]),
            }
        )
    return items


def _fetch_xray_analysis(connection, user_id: int, start_iso: str, end_iso: str) -> list[dict]:
    rows = connection.execute(
        """
        SELECT created_at, image_type, findings_json, abnormalities_json, severity, confidence,
               summary, recommendations_json, file_name
        FROM xray_analysis
        WHERE user_id = ? AND created_at BETWEEN ? AND ?
        ORDER BY created_at ASC, id ASC
        """,
        (user_id, start_iso, end_iso),
    ).fetchall()

    return [
        {
            "upload_date": _format_display_date(row["created_at"]),
            "body_part": row["image_type"],
            "findings": _safe_json_load(row["findings_json"], []),
            "severity": row["severity"],
            "ai_explanation": row["summary"],
            "suggested_action": _safe_json_load(row["recommendations_json"], []),
            "confidence": row["confidence"],
            "file_name": row["file_name"],
        }
        for row in rows
    ]


def _fetch_chat_history(connection, user_id: int, start_iso: str, end_iso: str) -> list[dict]:
    rows = connection.execute(
        """
        SELECT created_at, question, answer, topics_json, follow_up_json
        FROM chat_history
        WHERE user_id = ? AND created_at BETWEEN ? AND ?
        ORDER BY created_at ASC, id ASC
        """,
        (user_id, start_iso, end_iso),
    ).fetchall()

    return [
        {
            "created_at": _format_display_date(row["created_at"]),
            "question": row["question"],
            "answer": row["answer"],
            "topics_discussed": _safe_json_load(row["topics_json"], []),
            "follow_up_recommendations": _safe_json_load(row["follow_up_json"], []),
        }
        for row in rows
    ]


def _compute_statistics(
    connection,
    user_id: int,
    start_iso: str,
    end_iso: str,
    predictions: list[dict],
    xrays: list[dict],
    chats: list[dict],
) -> dict:
    last_activity = connection.execute(
        """
        SELECT created_at
        FROM user_activity
        WHERE user_id = ? AND created_at BETWEEN ? AND ?
        ORDER BY created_at DESC, id DESC
        LIMIT 1
        """,
        (user_id, start_iso, end_iso),
    ).fetchone()

    concern_counter = Counter()

    for prediction in predictions:
        concern_counter.update(_extract_keywords(prediction["predicted_disease"]))

    for xray in xrays:
        concern_counter.update(_extract_keywords(xray["body_part"]))

    for chat in chats:
        concern_counter.update(chat["topics_discussed"] or _extract_keywords(chat["question"]))

    most_common_health_concern = concern_counter.most_common(1)[0][0] if concern_counter else "General Health"

    return {
        "total_predictions": len(predictions),
        "total_xrays": len(xrays),
        "total_chats": len(chats),
        "most_common_health_concern": most_common_health_concern,
        "last_activity_date": _format_date_only(last_activity["created_at"]) if last_activity else None,
    }


def _build_executive_summary(
    user: dict,
    start_date: date,
    end_date: date,
    statistics: dict,
    concerns: list[str],
) -> dict:
    summary_lines = [
        f"This report summarizes the patient's activities from {start_date.strftime('%d %B %Y')} to {end_date.strftime('%d %B %Y')}.",
        "",
        "During this period:",
        f"• {statistics['total_predictions']} disease predictions were performed",
        f"• {statistics['total_xrays']} X-ray analyses were completed",
        f"• {statistics['total_chats']} chatbot consultations occurred",
    ]

    if concerns:
        summary_lines.extend(["", "Key concerns identified:"])
        summary_lines.extend([f"• {concern}" for concern in concerns[:3]])

    summary_lines.extend(["", "Recommended next actions:"])
    summary_lines.extend([f"• {action}" for action in _recommended_next_actions(concerns, statistics)])

    return {
        "executive_summary": "\n".join(summary_lines).strip(),
        "key_concerns": concerns[:3],
        "recommended_next_actions": _recommended_next_actions(concerns, statistics),
    }


def _recommended_next_actions(concerns: list[str], statistics: dict) -> list[str]:
    actions = [
        "Continue monitoring symptoms and keep a dated record of changes.",
        "Review the report with a qualified healthcare professional for formal interpretation.",
    ]

    if "Respiratory Issues" in concerns:
        actions.insert(0, "Follow up with a clinician if cough, fever, or breathing difficulty persists.")
    if "Musculoskeletal Concerns" in concerns:
        actions.insert(0, "Consider orthopedic review if pain, swelling, or movement limits continue.")
    if "Cardiovascular Concerns" in concerns:
        actions.insert(0, "Seek timely evaluation for chest discomfort, blood pressure spikes, or related symptoms.")

    if statistics.get("total_xrays", 0) > 0:
        actions.append("Bring the original imaging studies to follow-up visits when available.")

    return actions[:5]


def _concerns_from_results(predictions: list[dict], xrays: list[dict], chats: list[dict]) -> list[str]:
    counter = Counter()

    for prediction in predictions:
        counter.update(_extract_keywords(prediction["predicted_disease"]))
    for xray in xrays:
        counter.update(_extract_keywords(xray["body_part"]))
    for chat in chats:
        counter.update(chat["topics_discussed"] or _extract_keywords(chat["question"]))

    return [label for label, _ in counter.most_common(5)] or ["General Health"]


def generate_report_payload(user: dict, start_date: date, end_date: date) -> dict:
    start_iso, end_iso = _date_bounds(start_date, end_date)
    generated_at = utc_now()

    with db_session() as connection:
        sync_reporting_tables(connection, user["id"])
        predictions = _fetch_disease_predictions(connection, user["id"], start_iso, end_iso)
        xrays = _fetch_xray_analysis(connection, user["id"], start_iso, end_iso)
        chats = _fetch_chat_history(connection, user["id"], start_iso, end_iso)
        statistics = _compute_statistics(connection, user["id"], start_iso, end_iso, predictions, xrays, chats)
        concerns = _concerns_from_results(predictions, xrays, chats)
        summary = _build_executive_summary(user, start_date, end_date, statistics, concerns)
        recommendations = _recommended_next_actions(concerns, statistics)

        report_id = _generate_report_id(connection, generated_at)
        payload = {
            "report_id": report_id,
            "generated_at": generated_at,
            "start_date": start_date,
            "end_date": end_date,
            "user": {
                "id": user["id"],
                "first_name": user["first_name"],
                "last_name": user["last_name"],
                "email": user["email"],
                "phone": user.get("phone") or "",
                "created_at": user.get("created_at") or "",
            },
            "summary": summary,
            "disease_predictions": [
                {
                    key: value
                    for key, value in item.items()
                    if not key.startswith("_")
                }
                for item in predictions
            ],
            "xray_analyses": xrays,
            "chat_history": chats,
            "statistics": statistics,
            "recommendations": recommendations,
        }
        connection.execute(
            """
            INSERT INTO generated_reports (report_id, user_id, start_date, end_date, payload_json, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                report_id,
                user["id"],
                start_date.isoformat(),
                end_date.isoformat(),
                _safe_json_dump(payload),
                generated_at,
            ),
        )

    return payload


def _generate_report_id(connection, generated_at: str) -> str:
    date_token = generated_at[:10].replace("-", "")
    count = connection.execute(
        "SELECT COUNT(*) AS count FROM generated_reports WHERE created_at LIKE ?",
        (f"{generated_at[:10]}%",),
    ).fetchone()["count"]
    return f"RPT-{date_token}-{count + 1:03d}"


def get_saved_report(report_id: str, user_id: int) -> dict | None:
    with db_session() as connection:
        row = connection.execute(
            """
            SELECT payload_json
            FROM generated_reports
            WHERE report_id = ? AND user_id = ?
            """,
            (report_id, user_id),
        ).fetchone()
    if not row:
        return None
    return _safe_json_load(row["payload_json"], None)


def _load_logo_path() -> Path | None:
    for name in ["placeholder-logo.png", "icon-dark-32x32.png", "apple-icon.png"]:
        candidate = FRONTEND_DIR / name
        if candidate.exists():
            return candidate
    return None


def _report_sections(payload: dict) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []

    for item in payload.get("disease_predictions", []):
        sections.append(
            (
                f"Prediction - {item['predicted_disease']}",
                [
                    f"Date: {item['prediction_date']}",
                    f"Symptoms: {', '.join(item['input_symptoms']) or 'Not provided'}",
                    f"Predicted disease: {item['predicted_disease']}",
                    f"Confidence score: {item['confidence_score']}",
                    f"Risk level: {item['risk_level']}",
                    f"Recommendations: {', '.join(item['recommendations']) or 'Review with a clinician.'}",
                ],
            )
        )

    for item in payload.get("xray_analyses", []):
        sections.append(
            (
                f"X-Ray Analysis - {item['body_part']}",
                [
                    f"Upload date: {item['upload_date']}",
                    f"Findings: {', '.join(item['findings']) or 'None reported'}",
                    f"Severity: {item['severity']}",
                    f"AI explanation: {item['ai_explanation']}",
                    f"Suggested action: {', '.join(item['suggested_action']) or 'Review with a clinician.'}",
                ],
            )
        )

    for item in payload.get("chat_history", []):
        sections.append(
            (
                f"Chat - {item['created_at']}",
                [
                    f"Question: {item['question']}",
                    f"Answer: {item['answer']}",
                    f"Topics discussed: {', '.join(item['topics_discussed']) or 'General Health'}",
                    f"Follow-up: {', '.join(item['follow_up_recommendations']) or 'Monitor symptoms and follow up if needed.'}",
                ],
            )
        )

    return sections


def render_pdf_bytes(payload: dict) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            Image,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ModuleNotFoundError as exc:
        raise HTTPException(
            status_code=500,
            detail="PDF generation dependencies are missing. Install reportlab.",
        ) from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=54, bottomMargin=42)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitleCenter",
            parent=styles["Title"],
            alignment=TA_CENTER,
            textColor=colors.HexColor("#00D4FF"),
            fontSize=18,
            leading=22,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSection",
            parent=styles["Heading2"],
            textColor=colors.HexColor("#0B6BA8"),
            fontSize=12,
            leading=15,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportBody",
            parent=styles["BodyText"],
            leading=14,
            fontSize=9.5,
            textColor=colors.HexColor("#1E293B"),
        )
    )

    elements: list[Any] = []
    logo_path = _load_logo_path()
    if logo_path:
        elements.append(Image(str(logo_path), width=0.8 * inch, height=0.8 * inch))
    elements.append(Paragraph("MediVision AI Health Report", styles["ReportTitleCenter"]))
    elements.append(Spacer(1, 8))
    elements.append(
        Paragraph(
            "AI-generated medical-style summary. This report is not a medical diagnosis.",
            styles["ReportBody"],
        )
    )
    elements.append(Spacer(1, 12))

    metadata = [
        ["Report ID", payload["report_id"]],
        ["Generation Date", _format_display_date(payload["generated_at"])],
        ["Patient", f"{payload['user']['first_name']} {payload['user']['last_name']}"],
        ["Email", payload["user"]["email"]],
        ["Date Range", f"{_display_calendar_date(payload['start_date'])} - {_display_calendar_date(payload['end_date'])}"],
    ]
    metadata_table = Table(metadata, colWidths=[1.5 * inch, 4.9 * inch])
    metadata_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F4FF")),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#0F172A")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#2B8CFF")),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("LEADING", (0, 0), (-1, -1), 12),
            ]
        )
    )
    elements.append(metadata_table)
    elements.append(Spacer(1, 14))

    summary = payload["summary"]
    elements.append(Paragraph("Executive Summary", styles["ReportSection"]))
    elements.append(Paragraph(summary["executive_summary"].replace("\n", "<br/>"), styles["ReportBody"]))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("Activity Statistics", styles["ReportSection"]))
    stats_data = [
        ["Total Predictions", str(payload["statistics"]["total_predictions"])],
        ["Total X-Rays", str(payload["statistics"]["total_xrays"])],
        ["Total Chats", str(payload["statistics"]["total_chats"])],
        ["Most Common Concern", payload["statistics"]["most_common_health_concern"]],
        ["Last Activity Date", payload["statistics"]["last_activity_date"] or "N/A"],
    ]
    stats_table = Table(stats_data, colWidths=[2.2 * inch, 4.2 * inch])
    stats_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF5FF")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#0F172A")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ]
        )
    )
    elements.append(stats_table)
    elements.append(Spacer(1, 14))

    for heading, lines in _report_sections(payload):
        elements.append(Paragraph(heading, styles["ReportSection"]))
        for line in lines:
            elements.append(Paragraph(line, styles["ReportBody"]))
        elements.append(Spacer(1, 8))

    elements.append(Paragraph("Recommendations", styles["ReportSection"]))
    for item in payload.get("recommendations", []):
        elements.append(Paragraph(f"• {item}", styles["ReportBody"]))

    elements.append(Spacer(1, 10))
    elements.append(
        Paragraph(
            "Footer: MediVision AI patient summary report. Please consult a qualified healthcare professional for diagnosis or treatment decisions.",
            styles["ReportBody"],
        )
    )

    def draw_footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#2B8CFF"))
        canvas.setLineWidth(0.8)
        canvas.line(doc_obj.leftMargin, 28, doc_obj.pagesize[0] - doc_obj.rightMargin, 28)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawString(doc_obj.leftMargin, 16, "MediVision AI | AI-generated report")
        canvas.drawRightString(doc_obj.pagesize[0] - doc_obj.rightMargin, 16, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    doc.build(elements, onFirstPage=draw_footer, onLaterPages=draw_footer)
    return buffer.getvalue()


def render_docx_bytes(payload: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ModuleNotFoundError as exc:
        raise HTTPException(
            status_code=500,
            detail="DOCX generation dependencies are missing. Install python-docx.",
        ) from exc

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10)

    if logo_path := _load_logo_path():
        document.add_picture(str(logo_path), width=Inches(0.85))

    document.add_heading("MediVision AI Health Report", level=0)
    document.add_paragraph("AI-generated medical-style summary. This report is not a medical diagnosis.")

    meta = document.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    for label, value in [
        ("Report ID", payload["report_id"]),
        ("Generation Date", _format_display_date(payload["generated_at"])),
        ("Patient", f"{payload['user']['first_name']} {payload['user']['last_name']}"),
        ("Email", payload["user"]["email"]),
        ("Date Range", f"{_display_calendar_date(payload['start_date'])} - {_display_calendar_date(payload['end_date'])}"),
    ]:
        row = meta.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(payload["summary"]["executive_summary"])

    document.add_heading("Activity Statistics", level=1)
    stats = document.add_table(rows=0, cols=2)
    stats.style = "Table Grid"
    for label, value in [
        ("Total Predictions", payload["statistics"]["total_predictions"]),
        ("Total X-Rays", payload["statistics"]["total_xrays"]),
        ("Total Chats", payload["statistics"]["total_chats"]),
        ("Most Common Concern", payload["statistics"]["most_common_health_concern"]),
        ("Last Activity Date", payload["statistics"]["last_activity_date"] or "N/A"),
    ]:
        row = stats.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    def add_bullet_list(title: str, items: list[str]) -> None:
        document.add_heading(title, level=1)
        for item in items:
            document.add_paragraph(item, style="List Bullet")

    for item in payload.get("disease_predictions", []):
        document.add_heading(f"Prediction - {item['predicted_disease']}", level=2)
        document.add_paragraph(f"Date: {item['prediction_date']}")
        document.add_paragraph(f"Symptoms: {', '.join(item['input_symptoms']) or 'Not provided'}")
        document.add_paragraph(f"Predicted disease: {item['predicted_disease']}")
        document.add_paragraph(f"Confidence score: {item['confidence_score']}")
        document.add_paragraph(f"Risk level: {item['risk_level']}")
        add_bullet_list("Recommendations", item.get("recommendations", []))

    for item in payload.get("xray_analyses", []):
        document.add_heading(f"X-Ray Analysis - {item['body_part']}", level=2)
        document.add_paragraph(f"Upload date: {item['upload_date']}")
        document.add_paragraph(f"Findings: {', '.join(item['findings']) or 'None reported'}")
        document.add_paragraph(f"Severity: {item['severity']}")
        document.add_paragraph(f"AI explanation: {item['ai_explanation']}")
        add_bullet_list("Suggested action", item.get("suggested_action", []))

    for item in payload.get("chat_history", []):
        document.add_heading(f"Chat - {item['created_at']}", level=2)
        document.add_paragraph(f"Question: {item['question']}")
        document.add_paragraph(f"Answer: {item['answer']}")
        document.add_paragraph(f"Topics discussed: {', '.join(item['topics_discussed']) or 'General Health'}")
        add_bullet_list("Follow-up recommendations", item.get("follow_up_recommendations", []))

    add_bullet_list("Overall Recommendations", payload.get("recommendations", []))
    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.text = (
        "MediVision AI patient summary report. Please consult a qualified healthcare professional for diagnosis or treatment decisions."
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def load_generated_report(report_id: str, user_id: int) -> dict:
    payload = get_saved_report(report_id, user_id)
    if not payload:
        raise HTTPException(status_code=404, detail="Report not found.")
    return payload
